"""Generate SADDLE AdPulse Code Audit V2 as a Word Document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x3B, 0x3F, 0x51)
    h.font.name = 'Calibri'

BRAND = RGBColor(0x3B, 0x3F, 0x51)
RED = RGBColor(0xDC, 0x26, 0x26)
AMBER = RGBColor(0xD9, 0x77, 0x06)
GREEN = RGBColor(0x16, 0x65, 0x34)
GRAY = RGBColor(0x64, 0x74, 0x8B)


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Light Grid Accent 1'
    # Header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return t


def section_break():
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SADDLE ADPULSE')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = BRAND

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Comprehensive Code Audit & Test Plan')
run.font.size = Pt(18)
run.font.color.rgb = BRAND

for text in ['Version 2.0', 'January 29, 2026']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY

for _ in range(6):
    doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Prepared for Production Readiness Assessment')
run.font.size = Pt(11)
run.font.color.rgb = BRAND
run.italic = True

section_break()

# ════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    'Section 1: Executive Summary',
    'Section 2: Architecture Overview',
    'Section 3: Desktop Application Audit',
    'Section 4: Landing Site Audit',
    'Section 5: Archive & Codebase Hygiene',
    'Section 6: Mobile & Video Assessment',
    'Section 7: Security Review',
    'Section 8: Prioritized Remediation Plan',
    'Section 9: Appendix — File Health Matrix',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    p.paragraph_format.space_after = Pt(2)

section_break()

# ════════════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Section 1: Executive Summary', level=1)

doc.add_heading('What Changed Since V1.0 (Jan 22, 2026)', level=2)
doc.add_paragraph(
    'Since the initial audit, significant features have been added: auth/invitation system, '
    'impact dashboard, executive dashboard, mobile companion app, Remotion video generation, '
    'pricing/terms pages, and PostgreSQL migration. However, rapid iteration has accumulated '
    'substantial tech debt.'
)

add_table(
    ['Metric', 'V1.0 (Jan 22)', 'V2.0 (Jan 29)', 'Delta'],
    [
        ['Python LOC (desktop)', '~22,000', '~30,000+', '+36%'],
        ['HTML pages (landing)', '15', '21', '+40%'],
        ['CSS total (landing)', '~40KB', '~60KB', '+50%'],
        ['Dead/commented code', '~5%', '~15–20%', '+3x'],
        ['Archived directories', '1', '3 (10.2MB)', '+9x'],
        ['Platforms', '2', '4', '+2x'],
    ],
)

doc.add_heading('Top 5 Issues', level=2)
top5 = [
    ('Dashboard duplication', 'impact_dashboard.py (4,083 LOC) and executive_dashboard.py (1,926 LOC) share ~40% identical logic'),
    ('10.2MB of archived dead code', 'Sitting in the repo across 3 archive directories'),
    ('Landing site: 100% copy-paste', 'Nav/footer duplicated across all 21 HTML files (no templating)'),
    ('31% of script.js is commented-out', 'Dead code from disabled beta features'),
    ('44% of styles.css is unused', 'About page, modals, quiz styles loaded globally on every page'),
]
for i, (title, desc) in enumerate(top5, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title} — ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Severity Distribution', level=2)
add_table(
    ['Severity', 'Count', 'Description'],
    [
        ['CRITICAL', '4 issues', 'Blocking / high waste'],
        ['HIGH', '8 issues', 'Should fix soon'],
        ['MEDIUM', '12 issues', 'Fix when able'],
        ['LOW', '6 issues', 'Nice to have'],
    ],
)

section_break()

# ════════════════════════════════════════════════════════════════════════
# SECTION 2: ARCHITECTURE OVERVIEW
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Section 2: Architecture Overview', level=1)

doc.add_heading('Current Platform Map', level=2)
add_table(
    ['Platform', 'Technology', 'Size', 'Status'],
    [
        ['desktop/', 'Python / Streamlit', '30,000+ LOC', 'PRODUCTION'],
        ['landing/', 'HTML / CSS / JS', '21 pages', 'PRODUCTION'],
        ['mobile/', 'React Native / Expo', '34 TS files', 'DEVELOPMENT'],
        ['remotion-video/', 'React / Remotion', '16 scenes', 'COMPLETE'],
    ],
)

doc.add_heading('Desktop Module Breakdown', level=2)
add_table(
    ['Directory', 'Purpose', 'LOC'],
    [
        ['core/', 'DB, auth, data pipeline', '10,172'],
        ['features/', 'Business logic & UI tabs', '19,767'],
        ['ui/', 'Layouts, themes, components', '~3,000'],
        ['utils/', 'Helpers & formatters', '~500'],
        ['api/', 'External API clients', '~400'],
        ['config/', 'Settings & feature flags', '~300'],
    ],
)

doc.add_heading('Largest Modules (Risk Surface)', level=2)
add_table(
    ['Module', 'LOC', 'Risk', 'Reason'],
    [
        ['impact_dashboard.py', '4,083', 'HIGH', 'Too many responsibilities'],
        ['postgres_manager.py', '2,833', 'MEDIUM', 'Critical but focused'],
        ['optimizer.py', '2,662', 'LOW', 'Well-scoped'],
        ['assistant.py', '2,246', 'LOW', 'Isolated'],
        ['executive_dashboard.py', '1,926', 'HIGH', 'Redundant with impact_dashboard'],
        ['db_manager.py', '1,878', 'MEDIUM', 'Dual-mode complexity'],
        ['report_card.py', '1,817', 'LOW', 'Focused'],
        ['styles.css (landing)', '1,926 lines', 'HIGH', '44% unused'],
    ],
)

section_break()

# ════════════════════════════════════════════════════════════════════════
# SECTION 3: DESKTOP APPLICATION AUDIT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Section 3: Desktop Application Audit', level=1)

# 3.1
doc.add_heading('3.1 CRITICAL — Dashboard Module Duplication', level=2)
doc.add_paragraph(
    'Files: features/impact_dashboard.py + features/executive_dashboard.py'
)
doc.add_paragraph(
    'Both modules independently calculate revenue timelines with decision markers, '
    'KPI cards and metrics, ROAS/CVR/CPC calculations (identical formulas), and impact attribution logic. '
    'executive_dashboard.py imports from impact_dashboard.py but then re-implements similar logic, '
    'creating circular dependency risk.'
)
p = doc.add_paragraph()
r = p.add_run('Shared calculation pattern appearing in 4+ files:')
r.italic = True
doc.add_paragraph(
    "df['ROAS'] = (df['Sales'] / df['Spend']).replace([np.inf, -np.inf], 0).fillna(0)\n"
    "df['CVR']  = (df['Orders'] / df['Clicks'] * 100).replace([np.inf, -np.inf], 0).fillna(0)\n"
    "df['CPC']  = (df['Spend'] / df['Clicks']).replace([np.inf, -np.inf], 0).fillna(0)",
    style='No Spacing'
)
p = doc.add_paragraph()
r = p.add_run('Recommendation: ')
r.bold = True
p.add_run('Extract shared metric calculations into utils/calculations.py. Merge dashboard modules into a single module with view-mode parameter.')

# 3.2
doc.add_heading('3.2 HIGH — layout.py Code Quality Issues', level=2)
add_table(
    ['Line(s)', 'Issue', 'Severity'],
    [
        ['49–50', 'navigate_to() called twice consecutively (duplicate call)', 'HIGH'],
        ['625–677', 'Entire ROAS calculation + calc_efficiency() defined twice; first block is dead code', 'CRITICAL'],
        ['170–302', '130 lines of CSS embedded in Python strings', 'MEDIUM'],
        ['689–694', 'Comment says "silently handle" but displays error; redundant import st and pass', 'LOW'],
        ['122–123', 'Duplicate comment headers', 'LOW'],
    ],
)

# 3.3
doc.add_heading('3.3 HIGH — report_card.py Contradictory Logic', level=2)
doc.add_paragraph(
    'Lines 137–142 contain a calculation that overwrites itself:'
)
doc.add_paragraph(
    'Line 137: actual_roas = total_spend / total_sales  →  computes ACOS (wrong)\n'
    'Line 142: actual_roas = total_sales / total_spend  →  computes ROAS (correct)',
    style='No Spacing'
)
doc.add_paragraph(
    'Line 137 is dead code — immediately overwritten. Confused comments between the lines '
    'suggest this was debugged live and never cleaned up. Additionally, line 70 uses lowercase '
    'tuple type hint (-> tuple[bool, str]) which breaks Python <3.10 compatibility.'
)

# 3.4
doc.add_heading('3.4 MEDIUM — data_hub.py Incomplete Validation', level=2)
doc.add_paragraph(
    'Lines 341–347 contain a validation check that does nothing:'
)
doc.add_paragraph(
    'if not all(col in df_renamed.columns for col in required_cols):\n'
    '    # Fallback: Try to be lenient\n'
    '    pass  # NO ACTUAL LOGIC',
    style='No Spacing'
)
doc.add_paragraph(
    'Additional issues: session state initialization pattern repeated 3+ times (should be a helper), '
    'commented-out toast notification on line 280, inconsistent indentation.'
)

# 3.5
doc.add_heading('3.5 MEDIUM — ppcsuite_v4_ui_experiment.py Status', level=2)
doc.add_paragraph(
    'This 79KB file appears to be a deprecated entry point. Issues include: '
    'triple-duplicate comment blocks, dotenv soft-imported but used without checking import success, '
    'catches FileNotFoundError for st.secrets (should be KeyError), and comments about '
    '"delayed imports" immediately followed by eager imports.'
)
p = doc.add_paragraph()
r = p.add_run('Question: ')
r.bold = True
p.add_run('Is this file still the active entry point, or has it been superseded? If deprecated, archive it.')

# 3.6
doc.add_heading('3.6 MEDIUM — impact_dashboard.py Complexity', level=2)
doc.add_paragraph(
    'The _ensure_impact_columns() function performs 11 distinct operations in a single function. '
    'Should be decomposed into: _calculate_metrics(), _apply_guards(), _classify_market_tags(). '
    'Match type inference logic (infer_mt()) is also duplicated between this file and data_hub.py.'
)

# 3.7
doc.add_heading('3.7 LOW — theme.py Caching Inefficiency', level=2)
doc.add_paragraph(
    'get_cached_logo() uses @st.cache_data(ttl=3600) but theme changes call st.rerun() which '
    'invalidates all caches, making the TTL pointless. The 278 lines of nested CSS rules contain '
    'redundant selectors and hardcoded color values that bypass CSS variables already defined.'
)

section_break()

# ════════════════════════════════════════════════════════════════════════
# SECTION 4: LANDING SITE AUDIT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Section 4: Landing Site Audit', level=1)

# 4.1
doc.add_heading('4.1 CRITICAL — Navigation / Footer Duplication', level=2)
doc.add_paragraph(
    'Every HTML page contains identical copies of the navigation (~22 lines × 8+ copies = 176+ redundant lines) '
    'and footer (~46 lines × 8+ copies = 368+ redundant lines), plus identical head section content '
    '(font preconnects, meta tags, inline styles). Changing a nav link requires editing 8+ files.'
)
p = doc.add_paragraph()
r = p.add_run('Note: ')
r.bold = True
p.add_run('features.html is missing the JetBrains Mono font declaration present on other pages.')

# 4.2
doc.add_heading('4.2 CRITICAL — styles.css (1,926 lines) — 44% Waste', level=2)
add_table(
    ['Lines', 'Issue', 'Waste'],
    [
        ['67–76 vs 160–169', '.navbar defined twice with different background colors', '10 lines dead'],
        ['117–138 vs 197–213', '.primary-button-large defined twice, different box-shadows', '17 lines dead'],
        ['852–1150', 'About page styles in global stylesheet', '299 lines (unused on other pages)'],
        ['1152–1707', 'Quiz/audit modal + beta signup styles in global stylesheet', '555 lines (unused on most pages)'],
        ['313–386', 'Unused color utility classes (.highlight-green, .icon-green, etc.)', '~70 lines'],
    ],
)
doc.add_paragraph(
    'CSS Variable Duplication — Same color (#0F172A) mapped to 4 variable names: '
    '--color-slate-900, --color-wine, --color-dark, --color-text-primary. '
    'Unused variables: --shadow-sm, --color-wine-light, --color-dark-soft, --color-light-gray.'
)

# 4.3
doc.add_heading('4.3 CRITICAL — script.js — 31% Dead Code', level=2)
add_table(
    ['Lines', 'Description', 'Status'],
    [
        ['5–59', 'Persona state management (filtering, switching)', 'Commented out — "DISABLED FOR PRIVATE BETA"'],
        ['256–368', 'Pricing toggle + ICP card interactions', 'Commented out — "DISABLED FOR PRIVATE BETA"'],
        ['97', 'lastScroll variable assigned but never read', 'Dead variable'],
        ['223–226', '.hero-stats observer — element doesn\'t exist in HTML', 'Dead code'],
    ],
)
doc.add_paragraph(
    'Performance issues: Scroll event handler (lines 82–98) fires 60+ times/sec with no throttling; '
    'two separate ESC key listeners instead of one combined; Intersection Observer sets animation '
    'via inline styles (should use CSS classes); dynamic <style> tag injected via JS (lines 155–173) '
    'belongs in CSS file.'
)

# 4.4
doc.add_heading('4.4 HIGH — Page-Specific CSS Overlap', level=2)
add_table(
    ['File', 'Lines', 'Issue'],
    [
        ['features-styles.css', '316', 'Line 72: display: block !important — unnecessary !important'],
        ['agencies-styles.css', '377', 'Lines 281–341: .pricing-card duplicates rules from styles.css'],
        ['compare-styles.css', '346', 'Lines 9–29: heading styles could use global h1/h2'],
    ],
)

# 4.5
doc.add_heading('4.5 MEDIUM — audit.js Code Quality', level=2)
doc.add_paragraph(
    'Inline onclick handlers in dynamically generated HTML (line 91) — should use addEventListener. '
    'JSON.stringify inside template literal with hacky .replace(/\'/g, "&apos;") is fragile. '
    'Scoring multipliers 0.15 and 0.30 are magic numbers with no documentation. '
    'Fallback factor strings duplicated. Lines 298–319 reference HTML elements that don\'t exist (dead code).'
)

# 4.6
doc.add_heading('4.6 LOW — HTML Semantic Issues', level=2)
doc.add_paragraph(
    'No <main> element on any page (affects accessibility/SEO). '
    'Pricing cards use <div> instead of <article>. '
    'Missing <meta> description on some pages.'
)

section_break()

# ════════════════════════════════════════════════════════════════════════
# SECTION 5: ARCHIVE & CODEBASE HYGIENE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Section 5: Archive & Codebase Hygiene', level=1)

doc.add_heading('5.1 CRITICAL — Archive Directories (10.2MB in Repo)', level=2)
doc.add_paragraph(
    'These directories should not be in the repository. Git history preserves all prior states. '
    'Archives bloat clone times and confuse the active codebase.'
)
add_table(
    ['Directory', 'Size', 'Contents'],
    [
        ['desktop/dev_resources/_archived_20260114_161409/', '9.8MB',
         'Old landing pages, 25+ migration scripts, feature backups, Archive.zip'],
        ['desktop/dev_resources/_archived_cleanup_20260119/', '80KB',
         'Experimental backend, ingestion_v2 modules'],
        ['desktop/.backup_2024_12_18/', '1.5MB',
         'Manual point-in-time backup of 3 files'],
    ],
)

doc.add_heading('5.2 HIGH — Deprecated Files', level=2)
add_table(
    ['File', 'Size', 'Status'],
    [
        ['ppcsuite_v4_ui_experiment.py', '79KB', 'Labeled "experiment" — unclear if active'],
        ['desktop/dev_resources/frontend/', '~2MB', 'Abandoned Next.js experiment — no integration'],
        ['landing/_archive/', '56KB', 'Old beta pages + outdated documentation'],
        ['debug_sku_mapping.py', '~5KB', 'One-off debugging utility'],
        ['verify_uuid.py', '~2KB', 'One-off validation utility'],
        ['3× migrate_orphaned_accounts*.py', '~15KB', 'One-time migration scripts'],
    ],
)

doc.add_heading('5.3 MEDIUM — dev_resources Bloat', level=2)
doc.add_paragraph(
    'desktop/dev_resources/ contains 40+ test files (many one-off debugging scripts), '
    '20+ analysis scripts (one-time data explorations), 25+ completed migration scripts, '
    'and redesign artifacts / agent outputs / design briefs.'
)
p = doc.add_paragraph()
r = p.add_run('Recommendation: ')
r.bold = True
p.add_run('Keep migrations/ and documentation/. Move everything else to a separate archive repo or delete.')

section_break()

# ════════════════════════════════════════════════════════════════════════
# SECTION 6: MOBILE & VIDEO ASSESSMENT
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Section 6: Mobile & Video Assessment', level=1)

doc.add_heading('6.1 Mobile App (React Native / Expo)', level=2)
doc.add_paragraph(
    'Directory: mobile/ — 34 TypeScript files — Phase 6 (Polish)'
)
doc.add_paragraph(
    'Status: GREEN — Clean separation of screens, components, services, theme, and navigation. '
    'Proper TypeScript types throughout. No major code quality issues detected. '
    'Uses mock data (mockData.ts) — real API integration pending. Backend endpoint defined in mobile_api.py.'
)

doc.add_heading('6.2 Video Generation (Remotion)', level=2)
doc.add_paragraph(
    'Directory: remotion-video/ — 16 TSX scene files'
)
doc.add_paragraph(
    'Status: GREEN — Complete and well-organized. 10 primary scenes + 6 demo scenes. '
    '4K video generation capability. No issues detected.'
)

section_break()

# ════════════════════════════════════════════════════════════════════════
# SECTION 7: SECURITY REVIEW
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Section 7: Security Review', level=1)

doc.add_heading('7.1 Input Validation', level=2)
doc.add_paragraph(
    '• File uploads (data_hub.py lines 287–289): No file type validation — accepts any file extension\n'
    '• Session state: Multiple locations access st.session_state without initialization guards\n'
    '• Hardcoded currency: "AED" hardcoded in cache population instead of pulling from account settings'
)

doc.add_heading('7.2 Error Handling', level=2)
doc.add_paragraph(
    '• Silent pass statements in exception handlers (data_hub.py lines 347, 280) suppress errors without logging\n'
    '• ppcsuite_v4_ui_experiment.py catches FileNotFoundError instead of KeyError for st.secrets access'
)

doc.add_heading('7.3 Secrets Management', level=2)
doc.add_paragraph(
    '• .streamlit/secrets.toml exists in the repo tree — must be .gitignore\'d\n'
    '• python-dotenv used for env management — ensure .env is gitignored'
)

section_break()

# ════════════════════════════════════════════════════════════════════════
# SECTION 8: PRIORITIZED REMEDIATION PLAN
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Section 8: Prioritized Remediation Plan', level=1)

doc.add_heading('Phase 1: Immediate Cleanup (Day 1–2)', level=2)
add_table(
    ['#', 'Action', 'Impact', 'Effort'],
    [
        ['1', 'Delete 3 archive directories (10.2MB)', 'Repo hygiene', '15 min'],
        ['2', 'Remove or archive ppcsuite_v4_ui_experiment.py', 'Clarity', '15 min'],
        ['3', 'Delete commented-out code in script.js (167 lines)', 'Clean JS', '15 min'],
        ['4', 'Remove dead CSS rules in styles.css (duplicate .navbar, .primary-button-large)', 'Clean CSS', '30 min'],
        ['5', 'Fix report_card.py dead ACOS line and type hint', 'Correctness', '15 min'],
        ['6', 'Remove layout.py duplicate ROAS block (lines 625–650) and duplicate navigate_to() call', 'Correctness', '15 min'],
        ['7', 'Delete landing/_archive/', 'Repo hygiene', '5 min'],
    ],
)

doc.add_heading('Phase 2: Code Quality (Day 3–5)', level=2)
add_table(
    ['#', 'Action', 'Impact', 'Effort'],
    [
        ['8', 'Extract shared metric calculations (ROAS/CVR/CPC) into utils/calculations.py', 'DRY', '2 hrs'],
        ['9', 'Move about page CSS (299 lines) to about-styles.css', 'Performance', '1 hr'],
        ['10', 'Move modal/quiz CSS (555 lines) to modal-styles.css, lazy-load', 'Performance', '1 hr'],
        ['11', 'Consolidate CSS variables (4 names → 1 per color)', 'Maintainability', '1 hr'],
        ['12', 'Fix data_hub.py empty validation block (lines 341–347)', 'Correctness', '30 min'],
        ['13', 'Add scroll throttling to script.js navbar handler', 'Performance', '30 min'],
        ['14', 'Consolidate ESC key listeners in script.js', 'Clean code', '15 min'],
        ['15', 'Add missing JetBrains Mono font to features.html', 'Consistency', '5 min'],
    ],
)

doc.add_heading('Phase 3: Architecture (Week 2)', level=2)
add_table(
    ['#', 'Action', 'Impact', 'Effort'],
    [
        ['16', 'Merge executive_dashboard.py into impact_dashboard.py with view modes', 'Major DRY', '4–6 hrs'],
        ['17', 'Decompose _ensure_impact_columns() into 3 focused functions', 'Maintainability', '2 hrs'],
        ['18', 'Introduce HTML templating for landing nav/footer (Jinja2 or 11ty)', 'Major DRY', '4–6 hrs'],
        ['19', 'Move inline CSS from layout.py (130 lines) to theme system', 'Clean code', '2 hrs'],
        ['20', 'Replace audit.js inline onclick handlers with addEventListener', 'Clean code', '1 hr'],
        ['21', 'Unify match type inference (single source in utils/)', 'DRY', '1 hr'],
    ],
)

doc.add_heading('Phase 4: Ongoing', level=2)
add_table(
    ['#', 'Action', 'Impact'],
    [
        ['22', 'Replace silent pass handlers with proper logging', 'Debuggability'],
        ['23', 'Add file type validation to upload handlers', 'Security'],
        ['24', 'Ensure .env and secrets.toml are in .gitignore', 'Security'],
        ['25', 'Move dev_resources one-off scripts to archive repo', 'Hygiene'],
        ['26', 'Add <main> semantic elements to all HTML pages', 'Accessibility'],
    ],
)

section_break()

# ════════════════════════════════════════════════════════════════════════
# SECTION 9: APPENDIX — FILE HEALTH MATRIX
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('Section 9: Appendix — File Health Matrix', level=1)

doc.add_heading('Desktop Application Files', level=2)
add_table(
    ['File', 'LOC', 'Health', 'Key Issue'],
    [
        ['impact_dashboard.py', '4,083', 'RED', 'Too many responsibilities, 11-op function'],
        ['executive_dashboard.py', '1,926', 'RED', '~40% overlap with impact_dashboard'],
        ['layout.py', '~800', 'RED', 'Duplicate ROAS block, embedded CSS, dup nav call'],
        ['ppcsuite_v4_ui_experiment.py', '~2,000', 'RED', 'Potentially deprecated, wrong exception types'],
        ['report_card.py', '1,817', 'YELLOW', 'Dead ACOS line, type hint issue'],
        ['data_hub.py', '658', 'YELLOW', 'Empty validation, repeated patterns'],
        ['theme.py', '~300', 'YELLOW', 'Cache inefficiency, redundant CSS selectors'],
        ['postgres_manager.py', '2,833', 'GREEN', 'Heavy but necessary'],
        ['optimizer.py', '2,662', 'GREEN', 'Well-scoped'],
        ['optimizer_ui.py', '189', 'GREEN', 'Clean, well-structured (model to follow)'],
        ['assistant.py', '2,246', 'GREEN', 'Isolated'],
        ['db_manager.py', '1,878', 'GREEN', 'Dual-mode complexity acceptable'],
    ],
)

doc.add_heading('Landing Site Files', level=2)
add_table(
    ['File', 'Lines', 'Health', 'Key Issue'],
    [
        ['styles.css', '1,926', 'RED', '44% unused, duplicate rules, variable chaos'],
        ['script.js', '537', 'RED', '31% dead code, performance issues'],
        ['All HTML (21 files)', '~8,000', 'RED', 'No templating, full nav/footer duplication'],
        ['audit.js', '320', 'YELLOW', 'Magic numbers, inline handlers, dead upload code'],
        ['features-styles.css', '316', 'YELLOW', 'Unnecessary !important'],
        ['agencies-styles.css', '377', 'YELLOW', 'Duplicate .pricing-card rules'],
        ['how-styles.css', '266', 'GREEN', 'Minimal issues'],
        ['compare-styles.css', '346', 'GREEN', 'Minor heading overlap'],
    ],
)

doc.add_heading('Archive & Other', level=2)
add_table(
    ['Item', 'Size', 'Health', 'Status'],
    [
        ['_archived_20260114_161409/', '9.8MB', 'RED', 'Should not be in repo'],
        ['_archived_cleanup_20260119/', '80KB', 'RED', 'Should not be in repo'],
        ['.backup_2024_12_18/', '1.5MB', 'RED', 'Should not be in repo'],
        ['dev_resources/frontend/', '~2MB', 'RED', 'Abandoned experiment'],
        ['landing/_archive/', '56KB', 'RED', 'Outdated'],
        ['mobile/', '34 files', 'GREEN', 'Clean architecture'],
        ['remotion-video/', '16 scenes', 'GREEN', 'Complete, well-organized'],
    ],
)

# ── Footer ──────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Audit conducted January 29, 2026 — Version 2.0')
r.font.color.rgb = GRAY
r.font.size = Pt(9)
r.italic = True

# ── Save ────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), 'SADDLE_ADPULSE_CODE_AUDIT_V2.docx')
doc.save(out_path)
print(f"Saved: {out_path}")
